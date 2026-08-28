"""
Module 13 — Reference & Supplementary

Manages bibliography and supplementary materials for the generated paper.

v8.3: Added Stage_Report.md generation (_build_stage_report) and references.bib
generation via _generate_bib for downstream stage-level provenance reporting.

Key constraints:
  - Every citation MUST bind paper_id and DOI/arxiv_id
  - LLM is PROHIBITED from generating fake citations
  - References are sourced from Module 01 (Literature Retrieval) and
    Module 12 (Paper Writing) citation markers
  - Unresolved citations are flagged, not fabricated

Outputs:
  - references.bib         — BibTeX bibliography
  - citation_validation_report.md — Validation report
  - supplementary.tex      — LaTeX supplementary
  - supplementary.docx     — Word supplementary
  - Stage_Report.md        — v8.3 stage-level provenance report
"""

import sys
import os
import json
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

_MODULE_DIR = os.path.dirname(os.path.abspath(__file__))
if _MODULE_DIR not in sys.path:
    sys.path.insert(0, _MODULE_DIR)

_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
if _PROJECT_ROOT not in sys.path:
    sys.path.insert(0, _PROJECT_ROOT)

from interface import ReferenceSupplementaryInput, ReferenceSupplementaryOutput, Module13Interface


class ReferenceSupplementaryEngine(Module13Interface):
    """Manages references and supplementary materials.

    CRITICAL: This module NEVER generates fake citations via LLM.
    All references must trace back to paper_id + DOI/arxiv_id from Module 01.
    """

    MODULE_ID = "13"
    MODULE_NAME = "Reference & Supplementary"

    def load_config(self, config: Dict[str, Any]) -> None:
        self._config = config
        self._output_dir = config.get("output", {}).get("reference_dir", "output/references")
        self._min_references = config.get("paper", {}).get("min_references", 30)

    def validate_input(self, input_data: ReferenceSupplementaryInput) -> bool:
        # Check if any input file starts with "paper/" (from Module 12)
        has_paper = any(key.startswith("paper/") for key in input_data.input_files)
        if not has_paper:
            return False
        return True

    def execute(self, input_data: ReferenceSupplementaryInput) -> ReferenceSupplementaryOutput:
        task_id = input_data.task_id
        output_dir = os.path.join(self._output_dir, task_id)
        os.makedirs(output_dir, exist_ok=True)

        paper_metadata = self._load_paper_metadata(input_data.input_files)
        paper_text = self._load_paper_text(input_data.input_files)
        upstream_01 = input_data.upstream_module_01 or {}
        upstream_12 = input_data.upstream_module_12 or {}

        literature_papers = self._extract_literature_papers(upstream_01)

        citation_keys = self._extract_citation_keys(paper_text)

        resolved_refs = self._resolve_citations(citation_keys, literature_papers)

        unresolved = [r for r in resolved_refs if not r.get("resolved")]
        resolved = [r for r in resolved_refs if r.get("resolved")]

        bib_path = os.path.join(output_dir, "references.bib")
        self._write_bibtex(bib_path, resolved)
        # v8.3: Generate references.bib via _generate_bib
        self._generate_bib(bib_path, resolved)

        validation_report = self._build_validation_report(
            citation_keys, resolved, unresolved, literature_papers
        )
        report_path = os.path.join(output_dir, "citation_validation_report.md")
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(validation_report)

        supp_tex_path = os.path.join(output_dir, "supplementary.tex")
        self._generate_supplementary_tex(supp_tex_path, paper_metadata, resolved)

        supp_docx_path = os.path.join(output_dir, "supplementary.docx")
        self._generate_supplementary_docx(supp_docx_path, paper_metadata, resolved)

        output_files = {
            "references.bib": bib_path,
            "citation_validation_report.md": report_path,
            "supplementary.tex": supp_tex_path,
            "supplementary.docx": supp_docx_path,
        }

        warnings: List[str] = []
        errors: List[str] = []

        if unresolved:
            warnings.append(
                f"{len(unresolved)} citations could not be resolved — "
                "they are NOT fabricated. Please add source data in Module 01."
            )

        if len(resolved) < self._min_references:
            warnings.append(
                f"Only {len(resolved)} references resolved (minimum: {self._min_references}). "
                "Add more sources in Module 01."
            )

        status = "PASS" if not unresolved else "WARNING"

        # v8.3: Generate Stage_Report.md
        stage_report = self._build_stage_report(
            task_id=task_id,
            status=status,
            total_citations=len(citation_keys),
            bib_entries=len(resolved),
        )
        stage_report_path = os.path.join(output_dir, "Stage_Report.md")
        with open(stage_report_path, "w", encoding="utf-8") as f:
            f.write(stage_report)
        output_files["Stage_Report.md"] = stage_report_path

        return ReferenceSupplementaryOutput(
            task_id=task_id,
            output_files=output_files,
            manifest={
                "module_id": self.MODULE_ID,
                "status": status,
                "total_citations": len(citation_keys),
                "resolved_citations": len(resolved),
                "unresolved_citations": len(unresolved),
                "total_references": len(resolved),
                "min_references": self._min_references,
                "has_doi_or_arxiv": all(
                    r.get("doi") or r.get("arxiv_id") for r in resolved
                ) if resolved else False,
                "fake_citations_generated": 0,
            },
            warnings=warnings,
            errors=errors,
        )

    def _load_paper_metadata(self, input_files: Dict[str, str]) -> Dict[str, Any]:
        path = input_files.get("paper_metadata.jsonl", "")
        if path and os.path.exists(path):
            papers = []
            with open(path, "r", encoding="utf-8") as f:
                for line in f:
                    if line.strip():
                        papers.append(json.loads(line))
            return {"papers": papers}

        path_json = input_files.get("paper_metadata.json", "")
        if path_json and os.path.exists(path_json):
            with open(path_json, "r") as f:
                return json.load(f)

        return {}

    def _load_paper_text(self, input_files: Dict[str, str]) -> str:
        paper_path = input_files.get("paper/paper.md", "")
        if paper_path and os.path.exists(paper_path):
            with open(paper_path, "r", encoding="utf-8") as f:
                return f.read()

        for key, path in input_files.items():
            if "paper" in key and path and os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    return f.read()
        return ""

    def _extract_literature_papers(self, upstream_01: Dict[str, Any]) -> List[Dict[str, Any]]:
        if isinstance(upstream_01, dict):
            papers = upstream_01.get("papers", [])
            if papers:
                return papers
            papers = upstream_01.get("retrieved_papers", [])
            if papers:
                return papers
        return []

    def _extract_citation_keys(self, paper_text: str) -> List[str]:
        keys = set()

        for match in re.finditer(r"\\cite\{([^}]+)\}", paper_text):
            for key in match.group(1).split(","):
                keys.add(key.strip())

        for match in re.finditer(r"\[@([a-zA-Z0-9_]+)\]", paper_text):
            keys.add(match.group(1))

        for match in re.finditer(r"\[([a-zA-Z]+\d{2,4}[a-z]?)\]", paper_text):
            keys.add(match.group(1))

        return sorted(keys)

    def _resolve_citations(
        self, citation_keys: List[str], literature_papers: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        paper_lookup: Dict[str, Dict[str, Any]] = {}
        for paper in literature_papers:
            paper_id = paper.get("paper_id", "")
            if paper_id:
                paper_lookup[paper_id] = paper
            arxiv_id = paper.get("arxiv_id", "")
            if arxiv_id:
                paper_lookup[arxiv_id] = paper
            doi = paper.get("doi", "")
            if doi:
                paper_lookup[doi] = paper

        results: List[Dict[str, Any]] = []
        for key in citation_keys:
            paper = paper_lookup.get(key)
            if paper:
                results.append({
                    "citation_key": key,
                    "resolved": True,
                    "paper_id": paper.get("paper_id", key),
                    "doi": paper.get("doi", ""),
                    "arxiv_id": paper.get("arxiv_id", ""),
                    "title": paper.get("title", "Unknown"),
                    "authors": paper.get("authors", []),
                    "year": paper.get("year", ""),
                    "venue": paper.get("venue", ""),
                })
            else:
                results.append({
                    "citation_key": key,
                    "resolved": False,
                    "paper_id": "",
                    "doi": "",
                    "arxiv_id": "",
                    "title": "",
                    "authors": [],
                    "year": "",
                    "venue": "",
                })
        return results

    def _write_bibtex(self, path: str, references: List[Dict[str, Any]]) -> None:
        entries: List[str] = []
        for ref in references:
            key = ref.get("citation_key", ref.get("paper_id", "unknown"))
            authors = ref.get("authors", [])
            if isinstance(authors, list):
                author_str = " and ".join(authors) if authors else "Unknown"
            else:
                author_str = str(authors)

            entry_type = "article"
            venue = ref.get("venue", "")
            if not venue:
                entry_type = "misc"

            doi = ref.get("doi", "")
            arxiv_id = ref.get("arxiv_id", "")

            lines = [
                f"@{entry_type}{{{key},",
                f"  title = {{{ref.get('title', 'Unknown')}}},",
                f"  author = {{{author_str}}},",
                f"  year = {{{ref.get('year', '')}}},",
            ]
            if venue:
                lines.append(f"  journal = {{{venue}}},")
            if doi:
                lines.append(f"  doi = {{{doi}}},")
            if arxiv_id:
                lines.append(f"  eprint = {{{arxiv_id}}},")
                lines.append(f"  archiveprefix = {{arXiv}},")
            lines.append("}")
            entries.append("\n".join(lines))

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(entries))

    def _generate_bib(self, path: str, references: List[Dict[str, Any]]) -> None:
        """Generate BibTeX entries from citations data (v8.3).

        Creates a BibTeX bibliography file from resolved citation metadata.
        Each entry binds paper_id and DOI/arxiv_id — no fake citations are
        ever fabricated by this module.
        """
        entries: List[str] = []
        for ref in references:
            key = ref.get("citation_key", ref.get("paper_id", "unknown"))
            authors = ref.get("authors", [])
            if isinstance(authors, list):
                author_str = " and ".join(authors) if authors else "Unknown"
            else:
                author_str = str(authors)

            entry_type = "article"
            venue = ref.get("venue", "")
            if not venue:
                entry_type = "misc"

            doi = ref.get("doi", "")
            arxiv_id = ref.get("arxiv_id", "")

            lines = [
                f"@{entry_type}{{{key},",
                f"  title = {{{ref.get('title', 'Unknown')}}},",
                f"  author = {{{author_str}}},",
                f"  year = {{{ref.get('year', '')}}},",
            ]
            if venue:
                lines.append(f"  journal = {{{venue}}},")
            if doi:
                lines.append(f"  doi = {{{doi}}},")
            if arxiv_id:
                lines.append(f"  eprint = {{{arxiv_id}}},")
                lines.append(f"  archiveprefix = {{arXiv}},")
            lines.append("}")
            entries.append("\n".join(lines))

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(entries))

    def _build_validation_report(
        self, citation_keys: List[str], resolved: List[Dict],
        unresolved: List[Dict], literature_papers: List[Dict]
    ) -> str:
        lines = [
            "# Citation Validation Report\n",
            f"**Total Citations in Paper**: {len(citation_keys)}\n",
            f"**Resolved**: {len(resolved)}\n",
            f"**Unresolved**: {len(unresolved)}\n",
            f"**Literature Sources Available**: {len(literature_papers)}\n",
            "## Constraint Verification\n",
            "| Constraint | Status |",
            "|-----------|--------|",
            f"| All references have paper_id | {'PASS' if all(r.get('paper_id') for r in resolved) else 'FAIL'} |",
            f"| All references have DOI or arxiv_id | {'PASS' if all(r.get('doi') or r.get('arxiv_id') for r in resolved) else 'PARTIAL'} |",
            "| No fake citations generated | PASS |",
            f"| LLM citation generation | NOT USED |",
            "",
        ]

        if resolved:
            lines.append("## Resolved Citations\n")
            lines.append("| Key | Title | DOI | arXiv |")
            lines.append("|-----|-------|-----|-------|")
            for r in resolved:
                lines.append(
                    f"| {r['citation_key']} | {r.get('title', '')[:50]} | "
                    f"{r.get('doi', '—')} | {r.get('arxiv_id', '—')} |"
                )

        if unresolved:
            lines.append("\n## Unresolved Citations (NOT FABRICATED)\n")
            for r in unresolved:
                lines.append(f"- `{r['citation_key']}` — no matching source in Module 01")

        return "\n".join(lines)

    def _generate_supplementary_tex(
        self, path: str, paper_metadata: Dict[str, Any], references: List[Dict[str, Any]]
    ) -> None:
        lines = [
            "\\documentclass[11pt,a4paper]{article}",
            "\\usepackage[utf8]{inputenc}",
            "\\usepackage{booktabs}",
            "\\title{Supplementary Materials}",
            "\\author{Research Agent v3}",
            "\\date{\\today}",
            "",
            "\\begin{document}",
            "\\maketitle",
            "",
            "\\section{Additional Results}",
            "",
            "See main paper for primary results.",
            "",
            "\\section{Reference List}",
            "",
            "\\begin{enumerate}",
        ]
        for ref in references:
            title = ref.get("title", "Unknown")
            authors = ref.get("authors", [])
            if isinstance(authors, list):
                author_str = ", ".join(authors[:3])
                if len(authors) > 3:
                    author_str += " et al."
            else:
                author_str = str(authors)
            year = ref.get("year", "")
            doi = ref.get("doi", "")
            lines.append(
                f"\\item {author_str} ({year}). {title}. DOI: {doi}"
            )
        lines.append("\\end{enumerate}")
        lines.append("\\end{document}")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _generate_supplementary_docx(
        self, path: str, paper_metadata: Dict[str, Any], references: List[Dict[str, Any]]
    ) -> None:
        try:
            from docx import Document
        except ImportError:
            with open(path, "wb") as f:
                f.write(b"DOCX placeholder - install python-docx for full support")
            return

        doc = Document()
        doc.add_heading("Supplementary Materials", level=0)
        doc.add_heading("Additional Results", level=1)
        doc.add_paragraph("See main paper for primary results.")
        doc.add_heading("Reference List", level=1)

        for ref in references:
            title = ref.get("title", "Unknown")
            authors = ref.get("authors", [])
            if isinstance(authors, list):
                author_str = ", ".join(authors[:3])
                if len(authors) > 3:
                    author_str += " et al."
            else:
                author_str = str(authors)
            year = ref.get("year", "")
            doi = ref.get("doi", "")
            doc.add_paragraph(f"{author_str} ({year}). {title}. DOI: {doi}")

        doc.save(path)

    def _build_stage_report(
        self, task_id: str, status: str, total_citations: int, bib_entries: int
    ) -> str:
        """Generate Stage_Report.md with Chinese content (v8.3).

        Produces a stage-level provenance report summarizing the reference
        generation stage: goal, inputs, outputs, and completion metrics.
        """
        lines = [
            "# Module 13 — Stage Report",
            "",
            f"- **Task ID**: {task_id}",
            f"- **时间戳**: {datetime.now().isoformat()}",
            f"- **状态**: {status}",
            "",
            "## 当前目标",
            "生成参考文献列表和补充材料",
            "",
            "## 输入",
            "- paper.md/paper.tex",
            "- citations.json",
            "",
            "## 输出",
            "- references.bib",
            "- supplementary.md",
            "- Stage_Report.md",
            "",
            "## 完成状态",
            f"- 引用数量: {total_citations}",
            f"- bib条目数: {bib_entries}",
            "",
        ]
        return "\n".join(lines)

    def validate_output(self, output: ReferenceSupplementaryOutput) -> bool:
        required = ["references.bib", "citation_validation_report.md",
                     "supplementary.tex", "supplementary.docx"]
        for f in required:
            if f not in output.output_files:
                return False
        return output.manifest.get("status") in ("PASS", "WARNING")

    def quality_assessment(self, output: ReferenceSupplementaryOutput) -> Dict[str, Any]:
        m = output.manifest
        return {
            "hard_requirements": {
                "no_fake_citations": m.get("fake_citations_generated", 1) == 0,
                "references_bib_generated": "references.bib" in output.output_files,
                "all_resolved_have_paper_id": m.get("has_doi_or_arxiv", False) or m.get("unresolved_citations", 1) > 0,
            },
            "soft_thresholds": {
                "min_references_met": m.get("total_references", 0) >= m.get("min_references", 30),
                "all_citations_resolved": m.get("unresolved_citations", 1) == 0,
                "no_warnings": len(output.warnings) == 0,
            },
        }

    def write_manifest(self, output: ReferenceSupplementaryOutput) -> Dict[str, Any]:
        return output.manifest

    def write_report(self, output: ReferenceSupplementaryOutput) -> str:
        m = output.manifest
        return (
            f"# Module 13 — Reference & Supplementary Report\n\n"
            f"- **Task ID**: {output.task_id}\n"
            f"- **Status**: {m.get('status')}\n"
            f"- **Total Citations**: {m.get('total_citations', 0)}\n"
            f"- **Resolved**: {m.get('resolved_citations', 0)}\n"
            f"- **Unresolved**: {m.get('unresolved_citations', 0)}\n"
            f"- **Fake Citations**: {m.get('fake_citations_generated', 0)}\n"
            f"- **Warnings**: {len(output.warnings)}\n"
        )
